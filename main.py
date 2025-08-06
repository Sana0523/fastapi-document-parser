from fastapi import FastAPI, HTTPException, Request, Depends, Header
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field
from typing import List, Dict, Optional, Tuple, Any
import uvicorn
import json
import requests
import io
import time
import random
import os
from urllib.parse import urlparse
import hashlib
import logging
from contextlib import asynccontextmanager
import traceback
from datetime import datetime

# Document processing
from pypdf import PdfReader
import docx
import pdfplumber

# ML/AI components
import google.generativeai as genai
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np

# NLP components
import spacy
from nltk.corpus import stopwords
import nltk
import re

# Environment
from dotenv import load_dotenv

# Configure detailed logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

# Global variables for models
nlp = None
embedding_model = None
gemini_model = None

# Authorization setup
EXPECTED_TOKEN = "6613ee9a3bcb0925802224950bfad9d70f8be3907dc22442d035ae7798dbe14b"
security = HTTPBearer()

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Initialize models on startup with error handling"""
    global nlp, embedding_model, gemini_model
    
    try:
        logger.info("Starting model initialization...")
        
        # Download required NLTK data
        try:
            nltk.data.find('corpora/stopwords')
        except LookupError:
            logger.info("Downloading NLTK stopwords...")
            nltk.download('stopwords')
        
        # Initialize spaCy model
        try:
            logger.info("Loading spaCy model...")
            nlp = spacy.load("en_core_web_sm")
            logger.info("spaCy model loaded successfully")
        except OSError as e:
            logger.error("SpaCy model 'en_core_web_sm' not found. Install with: python -m spacy download en_core_web_sm")
            # Use a fallback or continue without spaCy
            nlp = None
        
        # Initialize sentence transformer
        logger.info("Loading sentence transformer...")
        embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        logger.info("Sentence transformer loaded successfully")
        
        # Configure Gemini
        GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
        if not GEMINI_API_KEY:
            logger.error("GEMINI_API_KEY not found in environment variables")
            raise ValueError("GEMINI_API_KEY not found in environment variables")
        
        logger.info("Configuring Gemini API...")
        genai.configure(api_key=GEMINI_API_KEY)
        
        safety_settings = [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
        ]
        
        gemini_model = genai.GenerativeModel('gemini-1.5-flash', safety_settings=safety_settings)
        logger.info("Gemini model configured successfully")
        
        logger.info("All models initialized successfully")
        yield
        
    except Exception as e:
        logger.error(f"Failed to initialize models: {str(e)}")
        logger.error(traceback.format_exc())
        raise
    
    logger.info("Shutting down...")

app = FastAPI(
    title="HackRX Document Query API", 
    version="1.0.0",
    description="Intelligent Document Q&A System for Insurance, Legal, HR, and Compliance",
    lifespan=lifespan
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Custom exception handler to fix the "dict object is not callable" error
@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    return JSONResponse(
        status_code=exc.status_code,
        content={"detail": exc.detail, "status_code": exc.status_code}
    )

@app.exception_handler(Exception)
async def general_exception_handler(request: Request, exc: Exception):
    logger.error(f"Unhandled exception: {str(exc)}")
    logger.error(traceback.format_exc())
    return JSONResponse(
        status_code=500,
        content={"detail": "Internal server error", "error": str(exc)}
    )

# Authentication dependency
async def verify_auth_token(credentials: Optional[HTTPAuthorizationCredentials] = Depends(security)):
    """Verify authorization token"""
    try:
        if credentials is None:
            logger.warning(f"Invalid or missing Authorization header: None")
            raise HTTPException(status_code=401, detail="Invalid or missing Authorization token")
        
        if credentials.credentials != EXPECTED_TOKEN:
            logger.warning(f"Invalid Authorization token: {credentials.credentials}")
            raise HTTPException(status_code=401, detail="Invalid or missing Authorization token")
        
        return True
    except Exception as e:
        logger.error(f"Unhandled exception: {str(e)}")
        raise HTTPException(status_code=401, detail="Invalid or missing Authorization token")

# Pydantic models
class HackRXRequest(BaseModel):
    documents: str = Field(..., description="URL to the document")
    questions: List[str] = Field(..., description="List of questions to answer")

class ClauseReference(BaseModel):
    chunk_id: str
    page_number: Optional[int]
    similarity_score: float
    text_snippet: str

class Answer(BaseModel):
    question: str
    answer: str
    confidence_score: float
    reasoning: str
    clause_references: List[ClauseReference]
    decision_rationale: Optional[str] = None

class HackRXResponse(BaseModel):
    answers: List[Answer]
    processing_time: float
    document_info: Dict[str, Any]

class DocumentChunk(BaseModel):
    text: str
    page_number: Optional[int] = None
    chunk_id: str
    embedding: Optional[List[float]] = None
    chunk_type: str = "general"  # general, clause, table, header

class DebugResponse(BaseModel):
    document_url: str
    text_length: int
    text_preview: str
    chunks_count: int
    first_chunk: Optional[str]
    page_count: int

# Global storage
document_cache = {}
vector_store = None
document_chunks = []

class DocumentProcessor:
    """Enhanced document processing with better error handling"""
    
    @staticmethod
    def download_document(url: str) -> bytes:
        """Download document with comprehensive error handling"""
        try:
            # Validate URL
            parsed_url = urlparse(url)
            if not all([parsed_url.scheme, parsed_url.netloc]):
                raise ValueError("Invalid URL format")
            
            # Create cache key
            url_hash = hashlib.md5(url.encode()).hexdigest()
            cache_file = f"cache_{url_hash}"
            
            # Check cache
            if os.path.exists(cache_file):
                logger.info(f"Using cached document: {cache_file}")
                with open(cache_file, 'rb') as f:
                    content = f.read()
                    if len(content) > 0:
                        return content
                    else:
                        os.remove(cache_file)  # Remove empty cache file
            
            # Download with proper headers and timeout
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'application/pdf,application/vnd.openxmlformats-officedocument.wordprocessingml.document,application/msword,*/*',
                'Accept-Language': 'en-US,en;q=0.9',
                'Accept-Encoding': 'gzip, deflate, br',
                'DNT': '1',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1',
            }
            
            logger.info(f"Downloading document from: {url}")
            
            with requests.Session() as session:
                response = session.get(
                    url, 
                    headers=headers, 
                    timeout=60,
                    stream=True,
                    allow_redirects=True
                )
                response.raise_for_status()
                
                # Read content with size limit (50MB)
                content = b''
                max_size = 50 * 1024 * 1024  # 50MB
                for chunk in response.iter_content(chunk_size=8192):
                    content += chunk
                    if len(content) > max_size:
                        raise ValueError("Document too large (>50MB)")
            
            if len(content) == 0:
                raise ValueError("Downloaded document is empty")
            
            # Cache the document
            try:
                with open(cache_file, 'wb') as f:
                    f.write(content)
                logger.info(f"Document cached: {cache_file}")
            except Exception as e:
                logger.warning(f"Failed to cache document: {str(e)}")
            
            logger.info(f"Document downloaded successfully, size: {len(content)} bytes")
            return content
            
        except requests.RequestException as e:
            logger.error(f"Network error downloading document: {str(e)}")
            raise HTTPException(500, f"Failed to download document: {str(e)}")
        except Exception as e:
            logger.error(f"Error downloading document: {str(e)}")
            raise HTTPException(500, f"Error processing document download: {str(e)}")
    
    @staticmethod
    def extract_text_from_pdf(content: bytes) -> Tuple[str, List[int]]:
        """Extract text from PDF with multiple fallback methods"""
        extracted_text = ""
        page_numbers = []
        
        try:
            logger.info("Attempting PDF extraction with pdfplumber")
            pdf_file = io.BytesIO(content)
            
            with pdfplumber.open(pdf_file) as pdf:
                total_pages = len(pdf.pages)
                logger.info(f"PDF has {total_pages} pages")
                
                for i, page in enumerate(pdf.pages):
                    try:
                        # Extract text
                        page_text = page.extract_text()
                        
                        # Extract tables if any
                        tables = page.extract_tables()
                        table_text = ""
                        if tables:
                            for table in tables:
                                for row in table:
                                    if row:
                                        table_text += " | ".join([cell or "" for cell in row]) + "\n"
                        
                        combined_text = (page_text or "") + "\n" + table_text
                        
                        if combined_text.strip():
                            extracted_text += f"\n--- Page {i+1} ---\n{combined_text}\n"
                            page_numbers.append(i + 1)
                            
                    except Exception as e:
                        logger.warning(f"Failed to extract page {i+1}: {str(e)}")
                        continue
                        
        except Exception as e1:
            logger.warning(f"pdfplumber failed: {str(e1)}, trying pypdf")
            
            try:
                pdf_file = io.BytesIO(content)
                pdf_reader = PdfReader(pdf_file)
                total_pages = len(pdf_reader.pages)
                logger.info(f"PDF has {total_pages} pages (pypdf)")
                
                for i, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            extracted_text += f"\n--- Page {i+1} ---\n{page_text}\n"
                            page_numbers.append(i + 1)
                    except Exception as e:
                        logger.warning(f"Failed to extract page {i+1} with pypdf: {str(e)}")
                        continue
                        
            except Exception as e2:
                logger.error(f"All PDF extraction methods failed: pdfplumber: {str(e1)}, pypdf: {str(e2)}")
                raise HTTPException(500, "Unable to extract text from PDF")
        
        if not extracted_text.strip():
            raise HTTPException(400, "No text could be extracted from the PDF")
        
        logger.info(f"Successfully extracted text from {len(page_numbers)} pages")
        return extracted_text.strip(), page_numbers
    
    @staticmethod
    def extract_text_from_docx(content: bytes) -> Tuple[str, List[int]]:
        """Extract text from DOCX with comprehensive content extraction"""
        try:
            logger.info("Extracting text from DOCX")
            docx_file = io.BytesIO(content)
            doc = docx.Document(docx_file)
            
            paragraphs = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        table_data.append(row_text)
                
                if table_data:
                    paragraphs.append("\n".join(table_data))
            
            # Extract headers/footers if possible
            try:
                for section in doc.sections:
                    if section.header:
                        header_text = "\n".join([p.text for p in section.header.paragraphs if p.text.strip()])
                        if header_text:
                            paragraphs.insert(0, f"HEADER: {header_text}")
                    
                    if section.footer:
                        footer_text = "\n".join([p.text for p in section.footer.paragraphs if p.text.strip()])
                        if footer_text:
                            paragraphs.append(f"FOOTER: {footer_text}")
            except Exception as e:
                logger.warning(f"Could not extract headers/footers: {str(e)}")
            
            if not paragraphs:
                raise HTTPException(400, "No text could be extracted from the DOCX")
            
            combined_text = "\n\n".join(paragraphs)
            logger.info(f"Successfully extracted {len(combined_text)} characters from DOCX")
            return combined_text, []  # DOCX doesn't have explicit page numbers
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise HTTPException(500, f"Failed to extract DOCX text: {str(e)}")
    
    @staticmethod
    def process_document(url: str) -> Tuple[str, List[int]]:
        """Main document processing with intelligent type detection"""
        try:
            content = DocumentProcessor.download_document(url)
            
            # Determine file type
            parsed_url = urlparse(url)
            file_path = parsed_url.path.lower()
            
            # Check file extension first
            if file_path.endswith('.pdf'):
                return DocumentProcessor.extract_text_from_pdf(content)
            elif file_path.endswith(('.docx', '.doc')):
                return DocumentProcessor.extract_text_from_docx(content)
            
            # Check content signature
            if content.startswith(b'%PDF'):
                return DocumentProcessor.extract_text_from_pdf(content)
            elif content.startswith(b'PK'):  # ZIP-based formats like DOCX
                return DocumentProcessor.extract_text_from_docx(content)
            
            # Try as plain text with different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text = content.decode(encoding)
                    if text.strip():
                        logger.info(f"Successfully decoded as {encoding}")
                        return text, []
                except UnicodeDecodeError:
                    continue
            
            # Final fallback: try PDF extraction
            try:
                return DocumentProcessor.extract_text_from_pdf(content)
            except:
                raise HTTPException(400, "Unable to determine document type or extract text")
                
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Document processing failed: {str(e)}")
            logger.error(traceback.format_exc())
            raise HTTPException(500, f"Document processing error: {str(e)}")

class TextChunker:
    """Advanced text chunking with semantic awareness"""
    
    @staticmethod
    def semantic_chunking(text: str, page_numbers: List[int], max_chunk_size: int = 600, overlap: int = 100) -> List[DocumentChunk]:
        """Create semantic chunks with improved logic for legal/insurance documents"""
        try:
            # Clean and normalize text
            text = re.sub(r'\n{3,}', '\n\n', text)
            text = re.sub(r'\s+', ' ', text)
            
            chunks = []
            chunk_id = 0
            
            # Split by major sections (clauses, articles, etc.)
            section_patterns = [
                r'\n(?=\s*(?:ARTICLE|SECTION|CLAUSE|CHAPTER)\s+[IVX\d]+)',
                r'\n(?=\s*\d+\.\s+[A-Z])',
                r'\n(?=\s*[A-Z][^a-z]{10,})',  # All caps headers
                r'\n(?=\s*---\s*Page\s+\d+\s*---)',  # Page breaks
            ]
            
            sections = [text]
            for pattern in section_patterns:
                new_sections = []
                for section in sections:
                    new_sections.extend(re.split(pattern, section))
                sections = new_sections
            
            # If no sections found, split by paragraphs
            if len(sections) <= 1:
                sections = text.split('\n\n')
            
            current_chunk = ""
            current_page = 1
            page_idx = 0
            
            for section in sections:
                section = section.strip()
                if not section:
                    continue
                
                # Update page number if we find a page marker
                page_match = re.search(r'---\s*Page\s+(\d+)\s*---', section)
                if page_match:
                    current_page = int(page_match.group(1))
                    page_idx = min(page_idx + 1, len(page_numbers) - 1) if page_numbers else page_idx
                    section = re.sub(r'---\s*Page\s+\d+\s*---', '', section).strip()
                    if not section:
                        continue
                
                # Determine chunk type
                chunk_type = TextChunker._determine_chunk_type(section)
                
                # Handle large sections by splitting into sentences
                if len(section) > max_chunk_size * 1.5:
                    sentences = re.split(r'(?<=[.!?])\s+', section)
                    for sentence in sentences:
                        if len(current_chunk) + len(sentence) > max_chunk_size and current_chunk:
                            # Create chunk
                            chunks.append(DocumentChunk(
                                text=current_chunk.strip(),
                                page_number=page_numbers[min(page_idx, len(page_numbers) - 1)] if page_numbers else current_page,
                                chunk_id=f"chunk_{chunk_id}",
                                chunk_type=chunk_type
                            ))
                            
                            # Create overlap
                            words = current_chunk.split()
                            overlap_words = min(overlap, len(words))
                            overlap_text = " ".join(words[-overlap_words:]) if overlap_words > 0 else ""
                            current_chunk = overlap_text + " " + sentence if overlap_text else sentence
                            chunk_id += 1
                        else:
                            current_chunk += " " + sentence if current_chunk else sentence
                else:
                    # Add entire section
                    if len(current_chunk) + len(section) > max_chunk_size and current_chunk:
                        chunks.append(DocumentChunk(
                            text=current_chunk.strip(),
                            page_number=page_numbers[min(page_idx, len(page_numbers) - 1)] if page_numbers else current_page,
                            chunk_id=f"chunk_{chunk_id}",
                            chunk_type=chunk_type
                        ))
                        
                        # Create overlap
                        words = current_chunk.split()
                        overlap_words = min(overlap, len(words))
                        overlap_text = " ".join(words[-overlap_words:]) if overlap_words > 0 else ""
                        current_chunk = overlap_text + "\n\n" + section if overlap_text else section
                        chunk_id += 1
                    else:
                        current_chunk += "\n\n" + section if current_chunk else section
            
            # Add final chunk
            if current_chunk.strip():
                chunks.append(DocumentChunk(
                    text=current_chunk.strip(),
                    page_number=page_numbers[min(page_idx, len(page_numbers) - 1)] if page_numbers else current_page,
                    chunk_id=f"chunk_{chunk_id}",
                    chunk_type=TextChunker._determine_chunk_type(current_chunk)
                ))
            
            logger.info(f"Created {len(chunks)} chunks from text")
            return chunks
            
        except Exception as e:
            logger.error(f"Text chunking failed: {str(e)}")
            logger.error(traceback.format_exc())
            raise HTTPException(500, f"Text chunking error: {str(e)}")
    
    @staticmethod
    def _determine_chunk_type(text: str) -> str:
        """Determine the type of content in a chunk"""
        text_lower = text.lower()
        
        if any(keyword in text_lower for keyword in ['clause', 'article', 'section']):
            return "clause"
        elif '|' in text and text.count('|') > 3:
            return "table"
        elif any(keyword in text_lower for keyword in ['header', 'title', 'chapter']):
            return "header"
        else:
            return "general"

class VectorStore:
    """Enhanced vector store with better search capabilities"""
    
    def __init__(self, dimension: int = 384):
        self.dimension = dimension
        self.index = faiss.IndexFlatIP(dimension)
        self.chunks = []
    
    def add_chunks(self, chunks: List[DocumentChunk]):
        """Add chunks to vector store with batch processing"""
        if not chunks:
            return
        
        try:
            logger.info(f"Encoding {len(chunks)} chunks...")
            texts = [chunk.text for chunk in chunks]
            
            # Process in batches to avoid memory issues
            batch_size = 16  # Reduced batch size for stability
            all_embeddings = []
            
            for i in range(0, len(texts), batch_size):
                batch_texts = texts[i:i + batch_size]
                logger.info(f"Processing batch {i//batch_size + 1}/{(len(texts) + batch_size - 1)//batch_size}")
                
                try:
                    batch_embeddings = embedding_model.encode(
                        batch_texts, 
                        normalize_embeddings=True,
                        show_progress_bar=False
                    )
                    all_embeddings.append(batch_embeddings)
                except Exception as e:
                    logger.error(f"Failed to encode batch {i//batch_size + 1}: {str(e)}")
                    # Create zero embeddings for failed batch
                    batch_embeddings = np.zeros((len(batch_texts), self.dimension))
                    all_embeddings.append(batch_embeddings)
            
            embeddings = np.vstack(all_embeddings)
            
            # Add to FAISS index
            self.index.add(embeddings.astype('float32'))
            
            # Store chunks with embeddings
            for i, chunk in enumerate(chunks):
                chunk.embedding = embeddings[i].tolist()
                self.chunks.append(chunk)
            
            logger.info(f"Successfully added {len(chunks)} chunks to vector store")
            
        except Exception as e:
            logger.error(f"Vector store addition failed: {str(e)}")
            logger.error(traceback.format_exc())
            raise HTTPException(500, f"Vector store error: {str(e)}")
    
    def search(self, query: str, k: int = 5, min_score: float = 0.1) -> List[Tuple[DocumentChunk, float]]:
        """Search for similar chunks with improved scoring"""
        if not self.chunks:
            return []
        
        try:
            query_embedding = embedding_model.encode([query], normalize_embeddings=True)
            scores, indices = self.index.search(
                query_embedding.astype('float32'), 
                min(k, len(self.chunks))
            )
            
            results = []
            for i, idx in enumerate(indices[0]):
                if idx >= 0 and idx < len(self.chunks):
                    score = float(scores[0][i])
                    if score > min_score:
                        results.append((self.chunks[idx], score))
            
            # Sort by score descending
            results.sort(key=lambda x: x[1], reverse=True)
            return results
            
        except Exception as e:
            logger.error(f"Vector search failed: {str(e)}")
            return []

class QueryProcessor:
    """Advanced query processing with improved LLM integration"""
    
    @staticmethod
    def _call_gemini_with_retry(prompt: str, max_retries: int = 3) -> str:
        """Call Gemini API with comprehensive retry logic"""
        for attempt in range(max_retries):
            try:
                response = gemini_model.generate_content(
                    prompt,
                    generation_config=genai.types.GenerationConfig(
                        temperature=0.1,
                        top_p=0.8,
                        top_k=40,
                        max_output_tokens=2048,
                    )
                )
                
                # Check for safety blocks
                if hasattr(response, 'candidates') and response.candidates:
                    candidate = response.candidates[0]
                    if hasattr(candidate, 'finish_reason'):
                        if candidate.finish_reason.name == "SAFETY":
                            raise ValueError("Response blocked by safety filters")
                
                if response.text and response.text.strip():
                    return response.text.strip()
                else:
                    raise ValueError("Empty response from Gemini")
                    
            except Exception as e:
                logger.warning(f"Gemini API attempt {attempt + 1} failed: {str(e)}")
                
                if attempt == max_retries - 1:
                    raise e
                
                # Exponential backoff with jitter
                wait_time = (2 ** attempt) + random.uniform(0, 1)
                time.sleep(wait_time)
        
        raise Exception("All Gemini API attempts failed")
    
    @staticmethod
    def analyze_query(query: str) -> Dict[str, Any]:
        """Enhanced query analysis with better domain understanding"""
        prompt = f"""
        Analyze this query for a legal/insurance document Q&A system. Return ONLY valid JSON.

        Query: "{query}"
        
        Analyze and return JSON with these fields:
        - "intent": Category (coverage, waiting_period, conditions, definition, benefit, exclusion, premium, claim, eligibility, procedure, general)
        - "key_entities": Important legal/insurance terms as strings
        - "question_type": Format (yes_no, explanation, list, definition, amount, procedure)
        - "search_keywords": Extended search terms including legal synonyms
        - "priority_terms": Most important 3-5 terms for search
        - "context_type": Domain (insurance, legal, hr, compliance, general)
        
        Example:
        {{
            "intent": "coverage",
            "key_entities": ["knee surgery", "orthopedic procedure"],
            "question_type": "yes_no", 
            "search_keywords": ["cover", "coverage", "knee", "surgery", "orthopedic", "procedure", "benefit", "eligible", "qualify"],
            "priority_terms": ["knee surgery", "coverage", "benefit"],
            "context_type": "insurance"
        }}
        
        Return only the JSON object.
        """
        
        try:
            response_text = QueryProcessor._call_gemini_with_retry(prompt)
            
            # Clean response
            response_text = response_text.strip()
            if response_text.startswith('```json'):
                response_text = response_text.replace('```json', '').replace('```', '').strip()
            elif response_text.startswith('```'):
                response_text = response_text.replace('```', '').strip()
            
            parsed_response = json.loads(response_text)
            
            # Validate required fields
            required_fields = ["intent", "key_entities", "question_type", "search_keywords"]
            for field in required_fields:
                if field not in parsed_response:
                    logger.warning(f"Missing field {field}, using fallback analysis")
                    return QueryProcessor._fallback_query_analysis(query)
            
            return parsed_response
            
        except Exception as e:
            logger.warning(f"Gemini query analysis failed: {str(e)}, using fallback")
            return QueryProcessor._fallback_query_analysis(query)
    
    @staticmethod
    def _fallback_query_analysis(query: str) -> Dict[str, Any]:
        """Comprehensive fallback analysis using NLP patterns"""
        query_lower = query.lower()
        
        # Enhanced intent patterns for legal/insurance domains
        intent_patterns = {
            "coverage": ["cover", "covered", "include", "benefit", "eligible", "qualify", "insure"],
            "waiting_period": ["waiting period", "wait", "how long", "when", "after", "before", "delay"],
            "conditions": ["condition", "requirement", "criteria", "must", "need", "qualify", "prerequisite"],
            "definition": ["what is", "define", "meaning", "definition", "explain", "describe"],
            "benefit": ["benefit", "amount", "limit", "maximum", "reimbursement", "payout", "compensation"],
            "exclusion": ["exclude", "not covered", "exception", "limitation", "restrict", "prohibited"],
            "premium": ["premium", "payment", "cost", "fee", "price", "monthly", "annual"],
            "claim": ["claim", "file", "submit", "process", "reimburse", "settlement"],
            "eligibility": ["eligible", "qualify", "qualification", "requirement", "criteria"],
            "procedure": ["procedure", "process", "steps", "how to", "method"]
        }
        
        detected_intent = "general"
        for intent, patterns in intent_patterns.items():
            if any(pattern in query_lower for pattern in patterns):
                detected_intent = intent
                break
        
        # Enhanced question type detection
        question_type = "explanation"
        if any(query_lower.startswith(word) for word in ["does", "is", "can", "will", "are", "do", "would"]):
            question_type = "yes_no"
        elif any(word in query_lower for word in ["what", "how", "why", "when", "where"]):
            if "what is" in query_lower or "define" in query_lower:
                question_type = "definition"
            elif "how much" in query_lower or "amount" in query_lower:
                question_type = "amount"
            elif "how to" in query_lower or "procedure" in query_lower:
                question_type = "procedure"
            else:
                question_type = "explanation"
        elif "list" in query_lower or "enumerate" in query_lower:
            question_type = "list"
        
        # Extract key entities using basic NLP
        key_entities = []
        # Remove common stop words and extract meaningful terms
        stop_words = set(stopwords.words('english')) if stopwords else set()
        words = re.findall(r'\b[A-Za-z]{3,}\b', query)
        key_entities = [word for word in words if word.lower() not in stop_words and len(word) > 2]
        
        # Generate search keywords
        search_keywords = key_entities.copy()
        
        # Add synonyms based on intent
        if detected_intent == "coverage":
            search_keywords.extend(["benefit", "insured", "policy", "plan"])
        elif detected_intent == "exclusion":
            search_keywords.extend(["not covered", "exception", "limitation"])
        elif detected_intent == "claim":
            search_keywords.extend(["reimbursement", "settlement", "payment"])
        
        # Context type detection
        context_type = "general"
        if any(word in query_lower for word in ["insurance", "policy", "coverage", "premium", "claim"]):
            context_type = "insurance"
        elif any(word in query_lower for word in ["legal", "law", "contract", "agreement", "clause"]):
            context_type = "legal"
        elif any(word in query_lower for word in ["employee", "hr", "benefits", "leave", "salary"]):
            context_type = "hr"
        elif any(word in query_lower for word in ["compliance", "regulation", "requirement", "standard"]):
            context_type = "compliance"
        
        return {
            "intent": detected_intent,
            "key_entities": key_entities[:10],  # Limit to top 10
            "question_type": question_type,
            "search_keywords": search_keywords[:15],  # Limit to top 15
            "priority_terms": key_entities[:5],  # Top 5 most important
            "context_type": context_type
        }
    
    @staticmethod
    def generate_answer(question: str, context_chunks: List[Tuple[DocumentChunk, float]], 
                       query_analysis: Dict[str, Any]) -> Answer:
        """Generate comprehensive answer using retrieved context"""
        if not context_chunks:
            return Answer(
                question=question,
                answer="I couldn't find relevant information in the document to answer this question.",
                confidence_score=0.0,
                reasoning="No relevant context found in the document",
                clause_references=[],
                decision_rationale="Insufficient information available"
            )
        
        # Prepare context for LLM
        context_text = ""
        clause_refs = []
        
        for i, (chunk, score) in enumerate(context_chunks[:5]):  # Top 5 chunks
            context_text += f"\n--- Context {i+1} (Score: {score:.3f}) ---\n{chunk.text}\n"
            clause_refs.append(ClauseReference(
                chunk_id=chunk.chunk_id,
                page_number=chunk.page_number,
                similarity_score=score,
                text_snippet=chunk.text[:200] + "..." if len(chunk.text) > 200 else chunk.text
            ))
        
        # Create comprehensive prompt
        prompt = f"""
        You are an expert document analyst specializing in {query_analysis.get('context_type', 'legal')} documents.
        
        Question: {question}
        
        Query Analysis:
        - Intent: {query_analysis.get('intent', 'general')}
        - Question Type: {query_analysis.get('question_type', 'explanation')}
        - Key Entities: {', '.join(query_analysis.get('key_entities', []))}
        
        Context from Document:
        {context_text}
        
        Instructions:
        1. Answer the question directly and comprehensively based on the provided context
        2. If the context doesn't fully answer the question, clearly state what information is missing
        3. For yes/no questions, provide a clear yes/no answer followed by explanation
        4. Quote specific clauses, sections, or phrases from the context when relevant
        5. If there are conflicting or unclear provisions, explain them
        6. Provide practical implications when appropriate
        
        Format your response as a detailed, professional answer suitable for legal/insurance contexts.
        """
        
        try:
            answer_text = QueryProcessor._call_gemini_with_retry(prompt)
            
            # Calculate confidence based on relevance scores and answer quality
            avg_score = sum(score for _, score in context_chunks) / len(context_chunks)
            confidence = min(0.95, avg_score * 1.2)  # Cap at 95%
            
            # Generate reasoning
            reasoning = f"Answer based on {len(context_chunks)} relevant document sections with average similarity score of {avg_score:.3f}. "
            reasoning += f"Primary context from {query_analysis.get('context_type', 'document')} domain focusing on {query_analysis.get('intent', 'general')} intent."
            
            return Answer(
                question=question,
                answer=answer_text,
                confidence_score=confidence,
                reasoning=reasoning,
                clause_references=clause_refs,
                decision_rationale=f"Based on semantic search and {query_analysis.get('question_type', 'general')} analysis"
            )
            
        except Exception as e:
            logger.error(f"Answer generation failed: {str(e)}")
            return Answer(
                question=question,
                answer="I encountered an error while processing your question. Please try rephrasing or contact support.",
                confidence_score=0.0,
                reasoning=f"Error in answer generation: {str(e)}",
                clause_references=clause_refs,
                decision_rationale="Processing error occurred"
            )

# API Endpoints
@app.get("/")
async def root():
    """Root endpoint with API information"""
    return {
        "message": "HackRX Document Query API",
        "version": "1.0.0",
        "status": "active",
        "endpoints": {
            "process": "/hackrx/run",
            "health": "/health",
            "docs": "/docs"
        }
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    global nlp, embedding_model, gemini_model
    
    status = {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "models": {
            "spacy": nlp is not None,
            "embedding": embedding_model is not None,
            "gemini": gemini_model is not None
        }
    }
    
    if not all(status["models"].values()):
        status["status"] = "degraded"
    
    return status

@app.post("/hackrx/run", response_model=HackRXResponse)
async def process_document_query(
    request: HackRXRequest,
    _: bool = Depends(verify_auth_token)
):
    """Main endpoint for document processing and question answering"""
    start_time = time.time()
    global vector_store, document_chunks
    
    try:
        logger.info(f"Processing request with {len(request.questions)} questions")
        
        # Step 1: Process document
        logger.info("Step 1: Processing document...")
        text, page_numbers = DocumentProcessor.process_document(request.documents)
        
        # Step 2: Create semantic chunks
        logger.info("Step 2: Creating semantic chunks...")
        chunks = TextChunker.semantic_chunking(text, page_numbers)
        document_chunks = chunks
        
        # Step 3: Build vector store
        logger.info("Step 3: Building vector store...")
        vector_store = VectorStore()
        vector_store.add_chunks(chunks)
        
        # Step 4: Process questions
        logger.info("Step 4: Processing questions...")
        answers = []
        
        for question in request.questions:
            logger.info(f"Processing question: {question}")
            
            # Analyze query
            query_analysis = QueryProcessor.analyze_query(question)
            
            # Search for relevant chunks
            search_results = vector_store.search(question, k=8)
            
            # Generate answer
            answer = QueryProcessor.generate_answer(question, search_results, query_analysis)
            answers.append(answer)
        
        # Calculate processing time
        processing_time = time.time() - start_time
        
        # Document info
        document_info = {
            "url": request.documents,
            "total_chunks": len(chunks),
            "total_pages": len(page_numbers) if page_numbers else 1,
            "text_length": len(text),
            "processing_method": "semantic_chunking_with_embeddings"
        }
        
        logger.info(f"Request completed successfully in {processing_time:.2f} seconds")
        
        return HackRXResponse(
            answers=answers,
            processing_time=processing_time,
            document_info=document_info
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Request processing failed: {str(e)}")
        logger.error(traceback.format_exc())
        raise HTTPException(500, f"Request processing error: {str(e)}")

@app.get("/debug/document", response_model=DebugResponse)
async def debug_document(
    url: str,
    _: bool = Depends(verify_auth_token)
):
    """Debug endpoint to inspect document processing"""
    try:
        text, page_numbers = DocumentProcessor.process_document(url)
        chunks = TextChunker.semantic_chunking(text, page_numbers)
        
        return DebugResponse(
            document_url=url,
            text_length=len(text),
            text_preview=text[:500] + "..." if len(text) > 500 else text,
            chunks_count=len(chunks),
            first_chunk=chunks[0].text if chunks else None,
            page_count=len(page_numbers) if page_numbers else 1
        )
        
    except Exception as e:
        logger.error(f"Debug processing failed: {str(e)}")
        raise HTTPException(500, f"Debug error: {str(e)}")

# if __name__ == "__main__":
#     uvicorn.run(
#         "main:app",
#         host="0.0.0.0",
#         port=8000,
#         reload=True,
#         log_level="info"
#     )